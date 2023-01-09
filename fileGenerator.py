import os
import pyautogui as auto
from PIL import ImageGrab
import time
from docx.shared import Inches
from docx import Document
from docx2pdf import convert

class FileGenerator:
    def __init__(self) -> None:
        self.document = Document()
        self.extension = ""
        self.fileList = []
        self.outputFileName = ""
        self.Map = {"py": self.runPythonCode,
                    "cpp": self.runCppCode, "java": self.runJavaCode}
        self.questionNum=1
        self.ignoreList=set()
        self.pdfFileName=""

    def getExtention(self) -> None:
        if self.extension=="":
            self.extension = input("Enter Files Extention: ")

    def clearScreen(self) -> None:
        os.system("cls")

    def generateFileList(self) -> None:
        res = True
        baseFileName = os.path.basename(__file__)
        self.ignoreList.add(baseFileName)

        self.fileList = [file for file in os.listdir() if file.endswith(self.extension) and file not in self.ignoreList]
        if self.extension not in self.Map:
            print(f"{self.extension} is not supported")
            res = False
        if len(self.fileList) == 0:
            print(f"{self.extension} files not found")
            res = False
        return res

    def getScreenshot(self) -> None:
        auto.hotkey("winleft", "shiftleft", "s")
        input()
        screenshot = ImageGrab.grabclipboard()
        screenshot.save("image.png")
        try:
            pass
        except Exception:
            print("An I/O Error Occoured at getScreenshot Please Restart The Process")

    def runPythonCode(self, fileName) -> None:
        self.clearScreen()
        os.system(f"python {fileName}")
        self.getScreenshot()

    def runCppCode(self, fileName) -> None:
        self.clearScreen()
        os.system(f"g++ {fileName} -o {fileName}.exe")
        os.system(f"{fileName}.exe")
        self.getScreenshot()

    def runJavaCode(self, fileName) -> None:
        nameWithoutExtention=fileName.removesuffix(".java")
        self.clearScreen()
        os.system(f"javac {fileName}")
        os.system(f"java {nameWithoutExtention}")
        self.getScreenshot()
    def runCode(self) -> None:
        pass



    def getQuestion(self, questionText) -> str:
        for index, char in enumerate(questionText.lower()):
            if 97 <= ord(char) <= 122:
                return questionText[index:]

    def generate(self) -> None:
        for _ in self.fileList:
            with open(_, "r") as file:
                question = self.getQuestion(file.readline())
                code = file.read()
                self.runCode(file.name)
                time.sleep(1)
                self.document.add_heading(f"{self.questionNum}. {question}")
                self.document.add_paragraph(code)
                self.document.add_picture(
                    "image.png", width=Inches(5.23), height=Inches(2.3))
                self.questionNum+=1
    def getOutputFileName(self):
        if self.outputFileName=="":
            self.outputFileName = input("Enter Output File Name: ")
        self.outputFileName += ".docx"
        self.pdfFileName=self.outputFileName.removesuffix(".docx")+".pdf"
    def start(self) -> None:
        self.getExtention()
        if not self.generateFileList():
            return
        self.runCode = self.Map[self.extension]
        self.generate()
        try:
            pass
        except Exception:
            print("An I/O Error Occoured at generate Please Restart The Process")
        time.sleep(1)
        os.remove("image.png")
        self.getOutputFileName()
        self.document.save(self.outputFileName)
        time.sleep(3)
        convert(self.outputFileName,self.pdfFileName)
        os.system(self.pdfFileName)

def main() -> None:
    object = FileGenerator()
    object.start()

if __name__ == "__main__":
    main()
