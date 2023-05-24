import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
import os
import platform
import pyautogui as auto
from PIL import ImageGrab
import time
from docx.shared import Inches
from docx import Document
from docx2pdf import convert


class FileGenerator:
    def __init__(self) -> None:
        self.document = Document()
        self.extension = ""
        self.fileList = []
        self.outputFileName = ""
        self.questionNum = 1
        self.ignoreList = set()
        self.pdfFileName = ""
        self.toMakePdf = True
        self.toMakeDocx = True
        self.workingDirectory = None
        self.os = ""
        self.Map = {"py": self.runPythonCode,
                    "cpp": self.runCppCode, "java": self.runJavaCode, "sh": self.runBashCode}
        self.clsMap = {"Windows": self.clearWindowsScreen,
                       "Linux": self.clearLinuxScreen}
        self.screenshotMap = {
            "Windows": self.getWindowsScreenshot, "Linux": self.getLinuxScreenshot}

    def getExtention(self) -> None:
        if self.extension == "":
            self.extension = input("Enter Files Extention: ")

    def clearWindowsScreen(self) -> None:
        os.system("cls")

    def clearLinuxScreen(self) -> None:
        os.system("clear")

    def clearScreen(self) -> None:
        pass

    def generateFileList(self) -> None:
        res = True
        baseFileName = os.path.basename(__file__)
        self.ignoreList.add(baseFileName)

        for file in os.listdir(self.workingDirectory):
            if file.endswith(self.extension) and file not in self.ignoreList:
                self.fileList.append(file)

        if len(self.fileList) == 0:
            print(f"{self.extension} files not found")
            self.showMessage(f"{self.extension} files not found")
            res = False
        return res

    def getWindowsScreenshot(self) -> None:
        auto.hotkey("winleft", "shiftleft", "s")
        input()
        screenshot = ImageGrab.grabclipboard()
        screenshot.save("image.png")

    def getLinuxScreenshot(self) -> None:
        os.system("rm image.png") if self.questionNum>1 else print(end="")
        os.system("scrot -s -f 'image.png'")
        input()
        # screenshot = ImageGrab.grabclipboard()
        # screenshot.save("image.png")

    def convertToPDF(self):
        if self.os == "Windows":
            convert(self.outputFileName, self.pdfFileName)
        elif self.os=="Linux":
            os.system(f"libreoffice --convert-to pdf {self.outputFileName}")

    def getScreenshot(self) -> None:
        pass

    def runPythonCode(self, fileName) -> None:
        self.clearScreen()
        os.system(f"python \"{fileName}\"")
        time.sleep(0.3)
        self.getScreenshot()

    def runCppCode(self, fileName) -> None:
        # os.system(f"g++ \"{fileName}\" -o \"{fileName}\".exe")
        # os.system(f"\"{fileName}\".exe")
        # self.getScreenshot()
        self.clearScreen()
        if self.os == "Windows":
            os.system(f"g++ \"{fileName}\" -o \"{fileName}\".exe")
            os.system(f"\"{fileName}\".exe")
        elif platform.system() == "Linux":
            os.system(f"g++ \"{fileName}\" -o \"{fileName}\".out")
            os.system(f"\"{fileName}\".out")
        time.sleep(0.3)
        self.getScreenshot()
    def runJavaCode(self, fileName) -> None:
        nameWithoutExtention = fileName.removesuffix(".java")
        self.clearScreen()
        os.system(f"javac \"{fileName}\"") if self.os=="Windows" else print(end="")
        os.system(f"java \"{fileName}\"")
        self.getScreenshot()

    def runBashCode(self, fileName):
        self.clearScreen()
        os.system(f"\"{fileName}\"")
        time.sleep(0.3)
        self.getScreenshot()

    def runCode(self) -> None:
        pass

    def showMessage(self, message):
        messagebox.showinfo("Pracfi", message)

    def getQuestion(self, questionText) -> str:
        for index, char in enumerate(questionText.lower()):
            if 97 <= ord(char) <= 122:
                return questionText[index:]

    def generate(self) -> None:
        for _ in self.fileList:
            currentFile = os.path.join(self.workingDirectory, _)
            with open(currentFile, "r") as file:
                question = self.getQuestion(file.readline())
                code = file.read()
                self.runCode(currentFile)
                time.sleep(1)
                self.document.add_heading(f"{self.questionNum}. {question}")
                self.document.add_paragraph(code)
                self.document.add_picture(
                    "image.png", width=Inches(5.23), height=Inches(2.3))
                self.questionNum += 1

    def getOutputFileName(self):
        if self.outputFileName == "":
            self.outputFileName = input("Enter Output File Name: ")
        self.outputFileName += ".docx"
        self.pdfFileName = self.outputFileName.removesuffix(".docx")+".pdf"

    def start(self) -> None:
        self.getExtention()
        if not self.generateFileList():
            return
        self.runCode = self.Map[self.extension]
        self.getScreenshot = self.screenshotMap[self.os]
        self.clearScreen = self.clsMap[self.os]
        self.generate()
        time.sleep(1)
        os.remove("image.png")
        self.getOutputFileName()
        self.clearScreen()
        print("Please Wait Processsing The Files")
        self.document.save(self.outputFileName)
        time.sleep(1)
        if self.toMakePdf == True:
            self.convertToPDF()
            os.system(self.pdfFileName) if self.os=="Windows" else os.system(f"open {self.pdfFileName}")
        if self.toMakeDocx == False:
            os.remove(self.outputFileName)
        self.clearScreen()


class Pracfi(tk.Tk):
    def __init__(self) -> None:
        super(Pracfi, self).__init__()
        self.title("Pracfi")
        self.minsize(400, 300)
        # self.maxsize(300, 300)
        self.startButton = tk.Button()
        self.outFileLabel = tk.Label()
        self.outFileInput = tk.Entry()
        self.folderSelectButton = tk.Button()

        # storage variables
        self.outFile = tk.StringVar(value="")  # contains output file name
        self.language = tk.StringVar(value=" ")  # contains selected language
        self.toMakePdf = tk.IntVar(value=True)
        self.toMakeDocx = tk.IntVar(value=True)
        self.workingDirPath = os.getcwd()
        self.isLinux = "disabled" if platform.system() != "Linux" else "normal"
        self.components()

    def start(self):
        self.mainloop()

    def handleStartButton(self):
        self.destroy()
        baseFileName = os.path.basename(__file__)
        object = FileGenerator()
        object.ignoreList.add(baseFileName)
        object.outputFileName = self.outFile.get()
        object.extension = self.language.get()
        object.toMakePdf = self.toMakePdf.get()
        object.toMakeDocx = self.toMakeDocx.get()
        object.workingDirectory = self.workingDirPath
        object.os = platform.system()
        object.start()

    def handleFolderSelectButton(self):
        self.workingDirPath = filedialog.askdirectory()

    def components(self):
        # Input Box
        self.outFileLabel.config(
            text="Enter Output File Name ", font=(20), pady=10, padx=10)
        self.outFileLabel.place(x=0, y=0)
        self.outFileInput.config(width=30, textvariable=self.outFile)
        self.outFileInput.place(y=13, x=200)

        self.startButton.place(x=100, y=240)
        self.startButton.config(
            text="Start Pracfi", command=self.handleStartButton, state="disabled")

        self.folderSelectButton.config(
            text="Select Folder ", command=self.handleFolderSelectButton, state="active")
        self.folderSelectButton.place(x=200, y=240)

        self.language.set(" ")
        python = tk.Radiobutton(self, text='Python',
                                variable=self.language, value='py', font=(20))
        cpp = tk.Radiobutton(
            self, text='C++', variable=self.language, value='cpp', font=(20))
        java = tk.Radiobutton(
            self, text='Java', variable=self.language, value='java', font=(20))
        bash = tk.Radiobutton(self, text='Bash', variable=self.language,
                              value='sh', font=(20), state=self.isLinux)
        python.place(x=0, y=50)
        cpp.place(x=0, y=80)
        java.place(x=0, y=110)
        bash.place(x=0, y=140)
        self.language.trace('w', self.onStateChange)
        self.outFile.trace('w', self.onStateChange)
        #  Check Boxes
        docxCheckbox = tk.Checkbutton(
            self, text="Word Document", variable=self.toMakeDocx, font=(20))
        pdfCheckbox = tk.Checkbutton(
            self, text="PDF", variable=self.toMakePdf, font=(20))
        docxCheckbox.place(x=0, y=180)
        pdfCheckbox.place(x=0, y=210)

    def onStateChange(self, *args):
        if (self.language.get() != " " and self.outFile.get() != ""):
            self.startButton.config(state='normal')
        else:
            self.startButton.config(state='disabled')


def main():
    obj = Pracfi()
    obj.start()


if __name__ == "__main__":
    main()
